#%pip install docx
#%pip install python-docx #this mutates docx? 
#%pip install pydantic
#%pip install mypy
# %pip install numpy
from typing import Optional, Dict, List, Any, Union, Tuple
from pydantic import BaseModel, root_validator, Field, ValidationError
import logging
from itertools import tee
import docx
from docx import Document
import numpy as np
from itertools import compress, chain
# import re
from collections import Counter
# import string



logger = logging.getLogger(__name__)

## Docx Pydantic Classes
### Helper Functions
def pairwise(iterable):
    '''mimic of itertools function not yet present in python 3.7.8
    pairwise('ABCDEFG') --> AB BC CD DE EF FG
    '''
    a, b = tee(iterable)
    next(b, None)
    return zip(a, b)
    
def logger_root_validation_error_messages(e, logger_details, suppress = [], run_enumeration: Optional[int] = None) -> Union[RuntimeError, TypeError]:      
   '''logger control for pydantic classes. Provides the ability to suppress certain validation errors, while still having the option to log them
   '''
   #TODO add ability to handle assertion errors
   if run_enumeration is not None:
      run_num = f"|run#{run_enumeration}|" #type: ignore 
   else:
      run_num = ""
   try:
      for err in e.errors():
         if err['type'] in suppress['type'] or err['msg'] in suppress['msg']:
            logger.info(f"|SUPRESSED|{logger_details['function']}|{type(e)}|para#{logger_details['paragraph_enumeration']}{run_num}, with validation? error: {err}")
            return TypeError("suppressed Validation Error")
         else:
            logger.error(f"|unsuppressed|{logger_details['function']}|{type(e)}|para#{logger_details['paragraph_enumeration']}{run_num}, with validation? error: {err}")
            return TypeError("un-suppressed Validation Error")
   except:
      logger.error(f"|unsuppressed|{logger_details['function']}|{type(e)}|para#{logger_details['paragraph_enumeration']}{run_num}, with error: {e}")
      return RuntimeError("non-validation error")
   return RuntimeError("non-validation error")

def str_strip_check(s:str) ->bool:#TODO, check if this is really needed?
   '''returns a boolean for if the input string would have characters removed by str.strip() 
   '''
   in_len = len(s)
   new_s = s.strip()
   out_len = len(new_s)
   purported_whitespace: bool = in_len != out_len
   return purported_whitespace

def closest(ranger, target): 
   '''#any target indeces occuring before the first ranger index will be ignored
   if ind target == ind ranger, that target will still be nested under that ranger
   '''
   if not isinstance(target,np.ndarray):
      target = np.array(target)
   for a,b in ranger:
      begin = np.searchsorted(target,a)
      end = np.searchsorted(target,b)
      _, out, target = np.split(target, [begin,end])
      yield list(out)
   yield list(target)

### Low Level "Base" Classes
class Docx_Paragraph (BaseModel):
   """input:   paragraph = your_paragraph_here
   when given a docx document's paragraph object, will parse it to a specified schema
   """
   # docx_document_paragraph: Optional[Any] #This should be validated below. Left optional because its inclusion causes problems with default repr and serialization
   para_text: str = Field(..., min_length = 1) ##required, must be string, must be 1 long or more
   para_first_line_indent: Optional[int] = Field(...) #Required, but must be int OR none. https://pydantic-docs.helpmanual.io/usage/models/#required-optional-fields 
   para_left_indent: Optional[int] = Field(...) #Required, but must be int OR none. https://pydantic-docs.helpmanual.io/usage/models/#required-optional-fields 

   class Config:
      validate_all = True
      # extra = 'forbid'
      validate_assignment = True
      smart_union = True
   
   @root_validator(pre=True) #TODO Try have post validator for runs only?
   def _docx_structure_check(cls, values: Dict[str, Any]) -> Dict[str, Any]:
      para = values.get("paragraph",False)
      assert isinstance(para, eval('docx.text.paragraph.Paragraph')), 'please enter a docx paragraph assigned to the variable "paragraph", in the form of     paragraph = your_paragraph_here'

      new_values: Dict[str, Any] = {}
      #extract para features, 
      new_values['para_text'] = para.text #type: ignore
      new_values['para_first_line_indent'] = para.paragraph_format.first_line_indent #type: ignore
      new_values['para_left_indent'] = para.paragraph_format.left_indent #type: ignore

      return new_values

class Docx_Run (BaseModel):
   """input:   run = your_run_here
   when given a docx document paragraphs run object, will parse it to a specified schema
   """
   run_text : str = Field(..., min_length = 1) #required, must be string, must be 1 long or more
   run_font_name : Optional[str] = Field(...) #required, must be string or None value
   run_font_size_pt : Optional[float] = Field(...)#Required, but must be float OR none value
   run_bold : Optional[bool] = Field(...) #Required, but must be bool OR none value
   run_italic : Optional[bool] = Field(...) #Required, but must be bool OR none value
   
   class Config:
      validate_all = True
      # extra = 'forbid'
      validate_assignment = True
      smart_union = True

   @root_validator(pre=True) #TODO Try have post validator for runs only?
   def _docx_structure_check(cls, values: Dict[str, Any]) -> Dict[str, Any]:
      run = values.get("run",False)
      assert isinstance(run, eval('docx.text.run.Run')), 'please enter a docx run assigned to the variable "run", in the form of     run = your_run_here'
      
      new_values : Dict[str, Any] = {}
      #loop through the runs in the paragraph and select the desired features
      new_values['run_text'] = run.text #type: ignore
      new_values['run_font_name'] = run.font.name #type: ignore
      if run.font.size is not None: #type: ignore
         new_values['run_font_size_pt'] = run.font.size.pt #type: ignore
      else: new_values['run_font_size_pt'] = None
      new_values['run_bold'] = run.bold #type: ignore
      new_values['run_italic'] = run.italic #type: ignore

      return new_values

class Docx_Run_List (BaseModel): #TODO refactor this to use run-aligned lists, so run obj can be used directly, and have its schema raised
   """input:   run_list = your_runs_in_a_list
   when given a list of docx document paragraphs run object, will parse it to a specified schema
   """
   #because the internals are validated, don't need to validate these other than that they were made into lists
   run_text : List[Any] = Field(...) #Required, must be list
   run_font_name : List[Any] = Field(...) #Required, must be list
   run_font_size_pt : List[Any] = Field(...) #Required, must be list
   run_bold : List[Any] = Field(...) #Required, must be list
   run_italic : List[Any] = Field(...) #Required, must be list

   @root_validator(pre=True) #TODO Try have post validator for runs only?
   def _docx_structure_check(cls, values: Dict[str, List[Any]]) -> Dict[str, Any]:
      from collections import defaultdict
      paragraph_enumeration = values.get('paragraph_enumeration',"<<FAILURE_paragraph_enumeration>>")
      runs = values.get("run_list",False)
      if not runs:
         raise ValueError('please enter a docx run list assigned to the variable "run_list", in the form of     run_list = your_run_list_here')
      new_values = defaultdict(list)
      suppress = {'type': ['value_error.any_str.min_length'], #ignore zero length run_text, per run validator
                  'msg': ['suppressed Validation Error']} #ignore suppressed errors earlier/lower in the stack      
      logger_details = {'function':'parsed_run', 'paragraph_enumeration':paragraph_enumeration }
      
      for run_enumumeration, run in enumerate(runs): #type: ignore
         try:
            parsed_run = Docx_Run(**{'run':run}) #this manner of root unpacking seems to give warnings since linter can't assess ahead of time
            assert isinstance(parsed_run, Docx_Run), 'RUNTIME_ERR - the docx run object did not return the type expected'
            for k,v in parsed_run.dict().items():
               new_values[k].append(v) 

         except BaseException as e:
            new_e = logger_root_validation_error_messages(e, logger_details, suppress,run_enumeration=run_enumumeration)
            raise new_e
             
      return new_values

### Aggregate Class with Internal Processing Functions
class Docx_Paragraph_and_Runs (BaseModel):
   """input:   paragraph = your_paragraph_here
   
   when given a docx document's paragraph object, will parse it to a specified schema
   """

   class Config:
      extra = 'allow'
      # arbitrary_types_allowed = True

   @root_validator(pre=True) #TODO Try have post validator for runs only?
   def _docx_structure_check(cls, values: Dict[str, Any]) -> Dict[str, Any]:
      if True: #reading/verifying inputs. Initializing needed variables/structures
         new_values: Dict[str, Any] = {}
         para = values.get("paragraph",False)
         assert isinstance(para, eval('docx.text.paragraph.Paragraph')), 'please enter a docx paragraph assigned to the variable "paragraph", in the form of     paragraph = your_paragraph_here'
         
         paragraph_enumeration: int = values.get('paragraph_enumeration',None)
         assert isinstance(paragraph_enumeration, int), "assertion error, bad paragraph count/paragraph_enumeration value passed. Please pass an integer"
         new_values['paragraph_enumeration'] = paragraph_enumeration

      #setting up error and logger handling
      #suppress these errors
      suppress = {'type': ['value_error.any_str.min_length'], #ignore zero length run_text, per run validator
                  'msg': ['suppressed Validation Error']} #ignore suppressed errors earlier/lower in the stack      

      #try to extract para features, 
      logger_details = {'function':'Docx_Paragraph', 'paragraph_enumeration':paragraph_enumeration }
      try: 
         parsed_paras = Docx_Paragraph(**{'paragraph':para}) #type: ignore
         for k,v in parsed_paras.dict().items():
            new_values[k] = v
      # except ValidationError as e:
      except BaseException as e:
         new_e = logger_root_validation_error_messages(e, logger_details, suppress)
         raise new_e

      #try to extract runs features
      logger_details = {'function':'Docx_Run_List', 'paragraph_enumeration':paragraph_enumeration }    
      try:
         parsed_runs = Docx_Run_List(**{'run_list':para.runs, 'paragraph_enumeration':paragraph_enumeration}) #type: ignore
         for k,v in parsed_runs.dict().items():
            new_values[k] = v
      except BaseException as e:
         new_e = logger_root_validation_error_messages(e, logger_details, suppress)
         raise new_e
         
      return new_values
      
   def interogate__para_text(self) -> str:
      t = getattr(self, 'para_text', "")
      # 
      if len(t) == 0:
         logger.warning('interogator did not find para_text')
      #    print("no para_text with:\n\t", self.dict())
      return t

   def get_run_text(self) -> Optional[List[str]]:
      run_texts = getattr(self,"run_text",None)
      if run_texts is not None:
         return run_texts
      else:
         logger.warning('interogator did not find run_text')
         return None

   def paragraph_logger(self,level:int,msg:str,print_bool:bool):
      if print_bool:
         print(msg)
      else:
         logger.log(level,msg)

   def single_run_feature_identify(self,params:Dict[str,Any]) -> Tuple[bool,Tuple[List[bool],List[Any]],Tuple[List[bool],List[Optional[str]]]]: 
      """if regex provided, must be in param dict with name 'text_regex_at_feature', and must be passed as a r'pattern' raw string
      return tuple of ('feature boolean', feature_Tuple[boolean mask, feature list], regex_tuple[boolean mask, regex match list])
      """
      if True: #reading/verifying inputs. Initializing needed variables/structures
         enumeration : Optional[int] = getattr(self,"paragraph_enumeration",None)
         assert isinstance(enumeration, int),f"bad value for 'paragraph_enumeration' {enumeration}"
         run_texts : Optional[List[str]] = getattr(self,'run_text',None)
         assert run_texts is not None, f"bad value for 'run_text' {self.__repr__()}"
         feature = params['docxFeature']
         assert isinstance(feature,str),f"bad value for parameter 'docxFeature'. Check params: {params}"
         text_regex_at_feature = params.get('text_regex_at_feature',False)
         regex_mask: List[bool] = []
         regex_matches: List[Optional[str]] = []
         position_requirement = params.get('position_requirement',False)
      values_from_runs: List[Optional[Union[float,bool]]] = getattr(self,feature,[None]) 
      value_mask: List[bool] = [True if x == params['value'] else False for x in values_from_runs]
      
      position_check = any(value_mask)
      if position_requirement is not False:
         if not isinstance(position_requirement,int):
            raise NotImplementedError('only single run positions are possible to enforce currently. Must pass an int for position index')
         try:
            position_check = value_mask[position_requirement]
         except:
            position_check=False
            pass

      if position_check:
         # print('text and value mask: ',run_texts,value_mask)
         # if text_regex_at_feature:
            # pattern = text_regex_at_feature
            # for text in run_texts:
            #    match = re.search(pattern, text) #type: ignore
            #    if match is not None:
            #       regex_mask.append(True)
            #       regex_matches.append(match.group(0))
            #       # print(repr(self))
            #    else:
            #       regex_mask.append(False)
            #       regex_matches.append(None)
            # print('regex and match: ',regex_mask,regex_matches)
            # # print(f'inside regex bool for para#{enumeration}\tregex_mask_is: {regex_mask}\t\tvalue_mask is: {value_mask}')
            # if not any(compress(value_mask,regex_mask)):
            #    return False, (value_mask, values_from_runs), (regex_mask, regex_matches) #does not have feature
         return True, (value_mask, run_texts), (regex_mask, regex_matches)  #has Feature
      else:
         return False, (value_mask, run_texts), (regex_mask, regex_matches) #does not have feature

   def modify_run_lists(self, drop_runs: Optional[List[int]] = None, add_runs: Optional[Tuple[int, List[List[Any]]]] = None, merge_runs : bool = False): #-> Optional[Dict[str, List[List[Any]]]]
      """given a list of indexes as 'drop' will drop those indexes from runlists, and return those dropped
      given a tuple with an integer index and list of lists (run aligned), will add those to entries to the runlists at that index
      given bool merge, will greedy merge all runs with the same run features EXCEPT run_text. Run_texts will be concatenated
      """
      if True: #reading/verifying inputs. Initializing needed variables/structures
         run_list_req_features: List[str] = Docx_Run_List.schema()['required']
         assert run_list_req_features[0] == 'run_text', "first feature in the schema should be run_text"
         para_enumeration = getattr(self, 'paragraph_enumeration',None)
         assert para_enumeration is not None, 'paragraph did not have an enumeration value'

         feature_run_lists : List[List[Any]] = []
         for f in run_list_req_features:
            feature_run_lists.append(getattr(self,f,[]))
         pivoted_run_lists = list(map(list, zip(*feature_run_lists)))
         number_of_runs : int = len(pivoted_run_lists)
         if number_of_runs < 1:
            raise ValueError('this paragraph does not have values in the run lists')
         merge_occured = False
         beginning_repr = self.__repr__()

      if drop_runs is not None:
         dropped_runs = {}
         num_dropped = 0
         for ind in drop_runs:
            mut_ind = ind - num_dropped #mutate pivot indexes as the pivot array is mutated
            dropped_runs[ind] = pivoted_run_lists.pop(mut_ind) #mutates pivoted_run_lists
            num_dropped +=1
         if number_of_runs == len(pivoted_run_lists):
            raise RuntimeError('the runs_lists were not shortened as expected')
         number_of_runs = len(pivoted_run_lists)
         # print(dropped_runs,pivoted_run_lists)
         feature_run_lists = list(map(list, zip(*pivoted_run_lists)))
         logger.info(f'para#{para_enumeration} had runs# {drop_runs} dropped. New run_text is: {feature_run_lists[0]}')

      if add_runs is not None:
         insert_ind : int = add_runs[0]
         add_lists = add_runs[1]
         assert len(add_lists[0]) == number_of_runs, "the added list of lists must have runs of the same length (feature space) as run_lists features in the schema: Docx_Run_List.schema()['required']"
         if insert_ind == -1:
            insert_ind = number_of_runs
         for lst in add_lists:
            pivoted_run_lists.insert(insert_ind,lst)
         number_of_runs = len(pivoted_run_lists)
         feature_run_lists = list(map(list, zip(*pivoted_run_lists)))

      if merge_runs is not False:
         i = 0
         still_merging = True
         # beginning_repr = self.__repr__()
         while still_merging:
            pairs = list(pairwise(list(range(len(pivoted_run_lists))))) #index pairs
            if len(pairs) < 1: #onely 1 run, which causes pairwise to yield empty lists since nothing to pair with
               break
            num_merged = 0
            for a,b in pairs: #where a,b are indexes in the pivoted run list (each index is one run)
               a -= num_merged #mutate pivot indexes after the pivot array has been mutated
               b -= num_merged
               if pivoted_run_lists[a][1:] == pivoted_run_lists[b][1:]: #if all features EXCEPT run_text are the same #TODO add ability to config which features to merge on
                  pivoted_run_lists[b][0] = pivoted_run_lists[a][0] + pivoted_run_lists[b][0]
                  pivoted_run_lists.pop(a)
                  num_merged +=1
                  merge_occured = True #flag for end of function, to determine if any changes need to be set to 'self'
               else: pass 
            if num_merged < 1: #if no merges where made in this iteration, merging is done. Else keep while loop since new merges may occur with new neighbors
               still_merging = False
         number_of_runs = len(pivoted_run_lists)
         feature_run_lists = list(map(list, zip(*pivoted_run_lists)))

      if any([drop_runs is not None, add_runs is not None, merge_occured]):
         for i, f in enumerate(run_list_req_features):
            self.__setattr__(f,feature_run_lists[i])

   def cleaner(self, execute_defaults: bool = True) -> bool : #params:Optional[Dict[str,Any]],
      """defaults to running "remove_para_leading_whitespace". This removes leading runs that are blank, and strips the first text run of any LEADING whitespace, if any is present.
      the params dict is not implemented currently
      returns bool value. True means cleaner would yield a valid para. False currently indicates all runs in para are whitespace.
      """
      if True: #reading/verifying inputs. Initializing needed variables/structures
         #TODO aggregate these getattrs so that every function doesn't need to get it themselves. Or simplify this with a function that has an assert bool to require it or not
         para_enumeration = getattr(self, 'paragraph_enumeration',None)
         assert para_enumeration is not None, 'paragraph did not have an enumeration value'

      def remove_para_leading_whitespace(start_ind : int = 0) -> bool: #run 
         run_text_list : List[str] = getattr(self, 'run_text',[''])
         num_runs = len(run_text_list)

         #control for all-whitespace paragraph
         para_text : Optional[str] = getattr(self, 'para_text',None)
         if isinstance(para_text,str):
            if len(para_text.strip()) == 0: #if para's text is ONLY whitespace
               logger.info(f'paragraph#{para_enumeration} with text "{para_text}" and runs {run_text_list} is all whitespace')
               return False
         else:
            raise RuntimeError(f'for paragraph#{para_enumeration}, the paragraph text attribute was not a string type')

         ind = start_ind
         droppable_runs : List[int] = [] #TODO this dropable section doesnt seem to be working correctly.
         while ind < num_runs:
            this_run_text = run_text_list[ind]
            stripped_run = this_run_text.lstrip() #TODO pass config to this to allow control of what can and can't be dropped.
            if len(stripped_run) == 0: #found ALL whitespace run. Need to iterate to see if next run is blank or has any leading whitespace
               droppable_runs.append(ind) #TODO convert this change to a an equivalent para_indent, since this paragraph likely has incorrect indents
               logger.info(f'paragraph#{para_enumeration} with text "{para_text}" had a run#{ind} with ONLY whitespace')
            elif len(stripped_run) < len(this_run_text): #found run that is NOT ALL whitespace, but had SOME. Will only happen once. Can stop now since this is the true beginning of this paragraph
               run_text_list[ind] = stripped_run
               self.__setattr__("run_text", run_text_list) #TODO convert this change to a an equivalent para_indent, since this paragraph likely has incorrect indents
               logger.info(f'paragraph#{para_enumeration} with text "{para_text}" had leading whitespace removed')
               break
            else: #Can stop now since this is the true beginning of this paragraph
               break
            ind +=1
            
         if len(droppable_runs) > 0: #if a whole run_text was whitespace only
            if len(droppable_runs) == num_runs: #if the whole paragraph was whitespace only
               raise RuntimeError(f'for paragraph#{para_enumeration}, all runs purported droppable whitespace, but para_text purported not')
            self.modify_run_lists(drop_runs = droppable_runs) #this removes whole runs, not just modifying the run_text.
            logger.info(f'paragraph#{para_enumeration} with text "{para_text}" tried to drop a run whitespace')
         return True

      if execute_defaults:
         cleaning_successful = remove_para_leading_whitespace()
         if not cleaning_successful:
            return False
         self.modify_run_lists(merge_runs = True)

      return True

### Draft of High Level Class
# class Fula_Entry (BaseModel): 
#    entity_word: List[str] #root, subroot, lemma
#    features: Optional[Dict[str,str]] = {} #contains features for this entity, ie: txt file features like location, POS, etc. Only applicable directly. Lemmas have POS, roots do not, etc
#    paragraphs_list: Dict[int,Any] #para enumeration, docx para obj
#    paragraphs_extr : List[Docx_Paragraph_and_Runs] #class defined above
#    sub_roots : List['Fula_Entry'] = [] #self reference
#    lemmas : List['Fula_Entry'] = [] #self reference


def read_docx(docx_filename:str , output_file: Optional[str] = None, verbose=False) -> Dict[str,list]:
   """takes a docx filepath, parses it with python-docx, then parses the docx paragraphs into Pydantic Dataclasses declared in this file
   'output_file' path may be passed to pickle the output instead of returning it in a dict of lists
   Has ability to somewhat control error handling of the pydantic classes, but this is not fully implemented as of 2022/07/30
   """
   #TODO add control logic for error tolerance/control
   document = Document(docx_filename)

   char_counts: Counter = Counter()
   docx_object_list = []
   parsed_object_list = []
   failed_paras_ind = [] #: List[Optional[Tuple[int,Docx_Paragraph_and_Runs,Any]]] = [] #: List[Optional[Tuple[int,Any,BaseException]]]= []
   handled_errors = []

   for i, para in enumerate(document.paragraphs):
      docx_object_list.append((i,para))
      try:
         entryObj = Docx_Paragraph_and_Runs(**{'paragraph': para, 'paragraph_enumeration': i})
         char_counts.update(entryObj.interogate__para_text())
         parsed_object_list.append((i,entryObj))
      except ValidationError as e:
         suppress = {
               # 'type': ['value_error.any_str.min_length'], #ignore zero length run_text, per run validator
               'msg': ['suppressed Validation Error'] #ignore suppressed errors earlier/lower in the stack      
         }
         for err in e.errors():
            if err['msg'] in suppress['msg']:
               handled_errors.append((i,para))
               pass
      except BaseException as e: #type: ignore
         failed_paras_ind.append((i,para))
         # raise e
   assert len(docx_object_list) == len(parsed_object_list) + len(handled_errors) + len(failed_paras_ind)

   if verbose: 
      print('total paras: ',len(docx_object_list))
      print('parsed paras: ',len(parsed_object_list))
      print('handled errors: ',len(handled_errors))
      print('failed paras: ',len(failed_paras_ind))
   #output logic
   outcomes_dict = {}         
   outcomes_dict['total_encountered_paragraphs'] = [len(document.paragraphs)]
   outcomes_dict['parsed_object_list'] = parsed_object_list
   outcomes_dict['handled_errors'] = handled_errors
   outcomes_dict['failed_paras_ind'] = failed_paras_ind
   outcomes_dict['char_counts'] = char_counts
   return outcomes_dict


def extract_features(entryObj: Docx_Paragraph_and_Runs, featureConfig: Dict[str,Any]) -> Union[Tuple[bool,str,List[bool],List[str]],Tuple[bool,bool,bool,bool]]:
   try:
      is_target, (target_mask, run_text), _ = entryObj.single_run_feature_identify(featureConfig)
      if is_target:
         text = ''.join(chain(compress(run_text,target_mask))).strip()
         #TODO control for usecase pulling text from multiple portions of the paragraph
         return (is_target, text, target_mask, run_text)
      else: 
         return (False,False,False,False)
   except Exception as e:
      raise e

if __name__ == '__main__':
   # import os
   # script_dir = os.path.dirname(__file__)
   # docx_filename = os.path.join(script_dir, "../test_data/pasted_docx page 1.docx")
   # # docx_filename = os.path.join(script_dir, "../test_data/Fula_Dictionary-repaired.docx")
   docx_filename = "test_data/pasted_docx page 1.docx"
   # docx_filename = "test_data/Fula_Dictionary-repaired.docx"
   parsed_to_dict = read_docx(docx_filename)
   print(len(parsed_to_dict))
   
